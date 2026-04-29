from __future__ import annotations

import logging
from pathlib import Path
from typing import List

import docx
from bs4 import BeautifulSoup

from epub_generator.config import SourceConfig
from epub_generator.ingestors.base import BaseIngestor, Metadata


logger = logging.getLogger(__name__)


class WordIngestor(BaseIngestor):
    """Parsear archivos .docx extraendo metadatos y capítulos."""

    def extract(self, source: SourceConfig, project_root: Path) -> tuple[Metadata, List[Chapter]]:
        """
        Extraer metadatos y capítulos de un archivo Word.

        Metadatos extraídos:
        - title: De la propiedad de título del documento o el nombre del archivo
        - author: Propiedad de autor del documento
        - language: Propiedad de idioma del documento

        Capítulos:
        - Se dividen por párrafos con estilo de título o por saltos de página grandes
        - Cada capítulo es un bloque de texto significativo
        """
        path = self.resolve_path(source, project_root)

        doc = docx.Document(str(path))

        # Extraer metadatos del documento
        title = self._extract_title(doc, path)
        author = self._extract_author(doc)
        language = self._extract_language(doc)

        metadata = Metadata(
            title=title,
            author=author,
            language=language,
        )

        # Extraer párrafos y dividir en capítulos
        paragraphs = self._extract_paragraphs(doc)

        # Dividir en capítulos basados en estilos o saltos de página
        chapters = self._group_into_chapters(paragraphs)

        return metadata, chapters

    def _extract_title(self, doc: docx.Document, path: Path) -> str:
        """Extraer el título del documento o usar el nombre del archivo."""
        # Intentar obtener el primer párrafo con estilo de título
        for para in doc.paragraphs:
            if para.style.name and "Title" in para.style.name or "Heading" in para.style.name:
                return para.text.strip()

        # Usar el nombre del archivo como título si no hay uno explícito
        name = path.stem
        return name.replace("-", " ").replace("_", " ").title()

    def _extract_author(self, doc: docx.Document) -> str | None:
        """Extraer el autor del documento."""
        # Intentar obtener el autor de las propiedades personalizadas
        custom_props = doc.core_properties or doc._doc.core_el or doc._doc.core_properties
        if hasattr(custom_props, "author"):
            author = getattr(custom_props, "author", None)
            if author:
                return author

        # Buscar propiedad de autor en propiedades personalizadas
        for prop in doc.core_properties or []:
            if hasattr(prop, "author"):
                return getattr(prop, "author", None)

        return None

    def _extract_language(self, doc: docx.Document) -> str | None:
        """Extraer el idioma del documento."""
        if doc.core_properties:
            lang = getattr(doc.core_properties, "lang", None)
            if lang:
                return lang
        return None

    def _extract_paragraphs(self, doc: docx.Document) -> List[str]:
        """Extraer todos los párrafos del documento como texto plano."""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                paragraphs.append(text)
        return paragraphs

    def _group_into_chapters(self, paragraphs: List[str]) -> List[Chapter]:
        """
        Agrupar párrafos en capítulos.

        Estrategia:
        - Un capítulo comienza con un párrafo de estilo de título/heading
        - O un capítulo comienza si hay más de N líneas desde el anterior
        """
        if not paragraphs:
            return []

        chapters = []
        current_chapter: List[str] = []
        chapter_title: str | None = None

        for para in paragraphs:
            # Verificar si el párrafo parece un título
            is_title = self._is_title_paragraph(para)

            if is_title:
                # Guardar capítulo anterior si existe
                if current_chapter:
                    title = chapter_title or self._generate_chapter_title(current_chapter)
                    chapters.append(Chapter(title=title, content="\n\n".join(current_chapter)))
                    current_chapter = []
                    chapter_title = para
                else:
                    chapter_title = para
            else:
                current_chapter.append(para)

        # Guardar el último capítulo
        if current_chapter:
            title = chapter_title or self._generate_chapter_title(current_chapter)
            chapters.append(Chapter(title=title, content="\n\n".join(current_chapter)))

        return chapters

    def _is_title_paragraph(self, text: str) -> bool:
        """Determinar si un texto parece ser un título de capítulo."""
        text_lower = text.lower()
        return text_lower.startswith("capitulo ") or \
               text_lower.startswith("capítulo ") or \
               text_lower.startswith("cap ") or \
               text_lower.startswith("section ") or \
               text_lower.startswith("sección ") or \
               len(text) < 50  # Títulos suelen ser cortos

    def _generate_chapter_title(self, content: List[str]) -> str:
        """Generar un título para un capítulo sin título explícito."""
        combined = " ".join(content).strip()
        return combined[:50].replace("\n", " ").title() or "Sin título"


class Chapter:
    """Representación de un capítulo extraído."""

    def __init__(self, title: str, content: str):
        self.title = title
        self.content = content

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "content": self.content,
        }
